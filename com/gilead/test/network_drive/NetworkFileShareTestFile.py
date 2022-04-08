'''
Created on Aug 20, 2019

@author: bkotamahanti

'''
import inspect
import os
import time
import shutil



class NetworkFileShareTestClass():
    def __init__(self, loggerRef):
        self.logger = loggerRef
        self.network_utils=NetworkUtils(loggerRef)
    
    
    def copy_files_between_drives_testcase_id_W10DA_001(self, src_drive, dest_drive, filename1, filename2=None ):
        try:
            func = inspect.currentframe().f_back.f_code
            self.logger.info("Starting execution of function {} :".format(func.co_name))
            source_file1=os.path.join(src_drive, filename1)
            dest_file1=os.path.join(dest_drive, filename1)
            if os.path.isfile(source_file1):
                exp_file_size1=self.network_utils.getFileSize(source_file1)
                start=time.time()
                shutil.copy2(source_file1, dest_drive)
                end=time.time()
                file1_time=end-start
                if filename2:
                    source_file2=os.path.join(src_drive, filename2)
                    dest_file2=os.path.join(dest_drive, filename2)
                    exp_file_size2=self.network_utils.getFileSize(source_file2)
                    start=time.time()
                    shutil.copy2(source_file2, dest_drive)
                    end=time.time()
                    file2_time=end-start
                self.logger.info("shutil copy function call done to destination {}".format(dest_drive))
                    
                assert self.network_utils.isFileExist(dest_file1) ,"File {} is not present in {}".format(source_file1, dest_drive)
                self.logger.info("Copied the source_file1 {} to drive {} ".format(source_file1, dest_drive))
                
                actual_file_size1=self.network_utils.getFileSize(dest_file1)
                
                assert exp_file_size1==actual_file_size1, "Expected File size {} KB didn't match with Actual file size {}".format(exp_file_size1, actual_file_size1)
                self.logger.info("Time taken to copy source_file1 {} of size {} KB to {} is {}ms".format(source_file1, str(exp_file_size1/1000), dest_drive, (file1_time)))
                
                if filename2:
                    assert self.network_utils.isFileExist(dest_file2) ,"File {} is not present in {}".format(source_file2, dest_drive)
                    self.logger.info("Copied the source_file2 {} to drive {} ".format(source_file2, dest_drive))
                    
                    actual_file_size2=self.network_utils.getFileSize(dest_file2)
                    assert exp_file_size2==actual_file_size2, "Expected File size {} KB didn't match with Actual file size {}".format(exp_file_size2, actual_file_size2)
                    self.logger.info("Time taken to copy source_file2 {} of size {} KB to {} is {}ms".format(source_file2, str(exp_file_size2/1000), dest_drive, (file2_time)))
                
        
            return True, "", func.co_name 
            
            
        except (AssertionError, AttributeError, TypeError) as e:
            import traceback
            traceback.print_exc()
            return False, e, func.co_name
    
    def copy_file_to_localdrive_modify_copyback_to_remotedrive_testcase_id_W10DA_002(self,src_drive, dest_drive, filename):
        try:
            func = inspect.currentframe().f_back.f_code
            self.logger.info("Starting execution of function {} :".format(func.co_name))
            if filename.endswith(".txt"):
                self.network_utils.killProcess("notepad.exe")
            elif filename.endswith(".docx"):
                self.network_utils.killProcess("WINWORD.EXE")
                
            source_file=os.path.join(src_drive, filename)
            dest_file=os.path.join(dest_drive, filename)
            if os.path.isfile(source_file):
                
                (initial_exp_file_size, start, end ) = self.network_utils.copyToDrive(source_file, dest_drive)
                self.logger.info("shutil copy function call done to destination {}".format(dest_drive))
                
                assert self.network_utils.isFileExist(dest_file) ,"File {} is not present in {}".format(source_file, dest_drive)
                self.logger.info("Copied the source_file {} to drive {} ".format(source_file, dest_drive))
                initial_actual_file_size=self.network_utils.getFileSize(dest_file)
                
                assert initial_exp_file_size==initial_actual_file_size, "Expected File size {} KB didn't match with Actual file size {}".format(initial_exp_file_size, initial_actual_file_size)
                self.logger.info("Time taken to copy source_file {} of size {} KB to {} is {}ms".format(source_file, str(initial_exp_file_size/1000), dest_drive, (end-start)))
                
            if dest_file.endswith(".txt"):
                status=self.network_utils.modifyTextFile(dest_file)
                assert status, "Unable to modify text file {}".format(dest_file)
                
            elif dest_file.endswith(".docx"):
                status=self.network_utils.modifyDocument(dest_file)
                assert status, "Unable to modify docx file {}".format(dest_file)

            self.logger.info("Destination File {} modified and ready to copy back to source drive {}".format(dest_file, source_file))
            (final_exp_file_size, start, end) = self.network_utils.copyToDrive(dest_file, src_drive)
            final_actual_file_size=self.network_utils.getFileSize(source_file)
            assert final_exp_file_size==final_actual_file_size, "Expected File size {} KB didn't match with Actual file size {}".format(initial_exp_file_size, initial_actual_file_size)
            self.logger.info("Time taken to copy source_file {} of size {} KB to {} is {}ms".format(source_file, str(initial_exp_file_size/1000), dest_drive, (end-start)))
            
                              
            return True, "", func.co_name 
            
            
        except (AssertionError, AttributeError, TypeError) as e:
            import traceback
            traceback.print_exc()
            return False, e, func.co_name
        
    
    
    
class NetworkUtils():
    def __init__(self,logRef):
        self.logger= logRef
    
    def isFileExist(self, file):
        return os.path.exists(file)
    
    def getFileSize(self, file):
        return os.path.getsize(file)
    
    def modifyTextFile(self, dest_file):
        with open(dest_file, 'r+') as txt_file_handle:
            for i in range(0,len(txt_file_handle.readlines())):
                txt_file_handle.write("File Modified....{}".format(i))
                txt_file_handle.write("\n")
        return True
    
    def copyToDrive(self, source_file, dest_drive):
        exp_file_size = self.getFileSize(source_file)
        start = time.time()
        shutil.copy2(source_file, dest_drive)
        end = time.time()
        return (exp_file_size, start, end)

    def modifyDocument(self, dest_file):
        from docx import Document;
        from io import BytesIO;
        with open(dest_file,"rb") as f:
            source_stream=BytesIO(f.read())
        
        document = Document(source_stream)
        source_stream.close()
        for i in range(30):      
            document.add_heading('Heading, level 1'+str(i), level=1)
            document.add_paragraph('Intense quote')
            document.add_paragraph('first item in unordered list')
            document.add_paragraph('first item in ordered list')
            
            records = (
                (3, '101', 'Spam'),
                (7, '422', 'Eggs'),
                (4, '631', 'Spam, spam, eggs, and spam')
                )

            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Qty'
            hdr_cells[1].text = 'Id'
            hdr_cells[2].text = 'Desc'
            for qty, num, desc in records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(qty)
                row_cells[1].text = num
                row_cells[2].text = desc

            document.add_page_break()
        
        document.save(dest_file)
        return True
    
    def isProcessExist(self,process_name):
        import psutil
        for proc in psutil.process_iter():
            # check whether the process_name name matches
            if proc.name() == process_name:
                self.logger.info("process "+str(process_name)+" exists")
                return True
        self.logger.info("process "+str(process_name)+" doesn't exist ")
        return False
    
    def sleepUntil(self, wait):
        self.logger.info("waiting for {} seconds for application to sync up or the element to show up".format(wait))
        time.sleep(wait)
    
    def killProcess(self, process_name):
        import psutil
        for proc in psutil.process_iter():
            # check whether the process_name name matches
            if proc.name() == process_name:
                self.logger.info("Found the process "+process_name+ " in process list and killing it now")
                proc.kill()
                
                if self.isProcessExist(process_name):
                    self.logger.info("waiting for some more time until the process is completely killed")
                    self.sleepUntil(5)
                self.logger.info("Killed the process "+process_name)