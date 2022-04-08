'''
Created on Oct 11, 2018
@author: pparamasivan
'''

from com.gilead.main.word.WordMainFile import WordMainClass
import inspect
from definitions import WORD_FILE_LOCATION
from selenium.common.exceptions import NoSuchElementException,\
    WebDriverException
from docx import Document


class WordTestClass():

    __file_location = WORD_FILE_LOCATION
        
    def __init__(self, utilsRef, loggerRef):
        self.utils = utilsRef
        self.logger = loggerRef
        self.word_maincode = WordMainClass(utilsRef, loggerRef)
        self.utils.removeFile(self.__file_location)
        
    
    '''TC ID: W10DA-44 '''        
    def word_open_save_close_testcase_id_W10DA_44(self):        
        try:
            func = inspect.currentframe().f_back.f_code
            self.logger.info("Starting execution of function {} :".format(func.co_name))

            '''start webdriver '''
            pid = self.utils.startWiniumDesktopDriver()
            self.logger.info("Launched the Winium Driver for Microsoft Word application. ")
            
            # TC Step 1
            self.logger.info(
                "Checking and killing the process WINWORD.EXE if any before executing the test case")
            self.word_maincode.killProcess()
              
            #TC Step 2,3,4,5,6,8
            
            self.word_maincode.createWordDocument(self.__file_location)
            self.utils.sleepUntil(3)
            assert self.utils.isFileExists(self.__file_location) , "Word Doc {} is not created ".format(self.__file_location)
            self.logger.info("Microsoft Word is created at {}".format(self.__file_location))
        
            
            self.utils.openFile(self.__file_location)
            self.utils.sleepUntil(3)
            assert self.word_maincode.isWordProcessExist(), "Microsoft Word WINWORD.EXE didn't start properly"            
            ''' Connect the already running Word process to winium Driver instead of opening new one.'''
            self.word_maincode.connectWordToWebDriver()
            self.logger.info("connected Word application to WebDriver")
            assert self.word_maincode.isWordElementExist(self.word_maincode.driver,**{'Name':'MsoDockTop'}), "Word Element is not present"
            self.logger.info("Checked the basic Word Doc elements present in Word GUI")          
            
            content_text = self.word_maincode.getWordDocumentContentText()
            self.logger.info("Doc content is ====> " + content_text)
            
            assert 'I love Python' in content_text, "Word doc content is not present"
            self.logger.info("Verified the Word Document Content text")
            
            #TC Step 9
            self.word_maincode.clickCloseButtonOfWordApp()
            self.utils.sleepUntil(3)
               
            #TC Step 10
            self.word_maincode.isWordProcessExist() == False, "Microsoft Word process still running"
            return True, "", func.co_name

        except (AssertionError, AttributeError, TypeError, WebDriverException, NoSuchElementException) as e:
            return False, e, func.co_name
        finally:
            '''kill the web driver '''
            pid.terminate()
            self.logger.info("Killed the webDriver process of Word application")
            self.word_maincode.killProcess()
 
    '''TC ID: W10DA-45 '''        
    def word_create_document_donot_save_testcase_id_W10DA_045(self):        
        try:
            func = inspect.currentframe().f_back.f_code
            self.logger.info("Starting execution of function {} :".format(func.co_name))

            '''start webdriver '''
            pid = self.utils.startWiniumDesktopDriver()
            self.logger.info("Launched the Winium Driver for Microsoft Word application. ")
            
            # TC Step 1 Check the process WINWORD.EXE not exist in task manager
            self.logger.info(
                "Checking and killing the process WINWORD.EXE if any before executing the test case")
            self.word_maincode.killProcess()

            #TC Step 2 Open MS Word
            self.word_maincode.startMicroSoftWordApp("WINWORD.EXE")
            self.utils.sleepUntil(5)

            assert self.word_maincode.isWordProcessExist() , "Microsoft Word process didn't start"
            self.logger.info("Microsoft Word process started")
            
            self.word_maincode.connectWordToWebDriver()
            self.logger.info("connected Microsoft Word app to WebDriver ")
            
            #TC Step 3  Click Blank Document 
            self.word_maincode.clickBlankDocument()
            self.logger.info("Clicked Blank Document layout")
            
            #TC Step 4  Enter multiple lines of text
            self.word_maincode.enterTextInDocument("My Word document!")
            self.logger.info("Entered 'My Word document!' in text area")
            self.utils.sleepUntil(3)
            
            #TC Step 5 Select Menu > File > Close (do not save)
            self.word_maincode.clickWordDontSaveMenuOption()
            self.logger.info("Clicked Close button of Word application")

            #TC Step 6 Click Don't Save 
            self.word_maincode.clickDontSaveInMsPopUp()
            self.logger.info("Clicked Dont save in MSPop Up")
            self.utils.sleepUntil(3)
               
            #TC Step 7  Check in Task Manager if Word process is running
            self.word_maincode.isWordProcessExist() == False, "Microsoft Word process still running"
            return True, "", func.co_name
        
        except (AssertionError, AttributeError, TypeError, WebDriverException, NoSuchElementException) as e:
            return False, e, func.co_name
        finally:
            '''kill the web driver '''
            pid.terminate()
            self.logger.info("Killed the webDriver process of Word application")
            self.word_maincode.killProcess()
 

    '''TC ID: W10DA-46 '''                    
    def word_delete_file_validate_error_testcase_id_W10DA_46(self):        
        try:
            func = inspect.currentframe().f_back.f_code
            self.logger.info("Starting execution of function {} :".format(func.co_name))

            '''start webdriver '''
            pid = self.utils.startWiniumDesktopDriver()
            self.logger.info("Launched the Winium Driver for Microsoft Word application. ")
            
            # TC Step 1
            self.logger.info(
                "Checking and killing the process WINWORD.EXE if any before executing the test case")
            self.word_maincode.killProcess()
              
            #TC Step 2,3,4
            
            self.word_maincode.createWordDocument(self.__file_location)
            self.utils.sleepUntil(3)
            assert self.utils.isFileExists(self.__file_location) , "Word Doc {} is not created ".format(self.__file_location)
            self.logger.info("Microsoft Word is created at {}".format(self.__file_location))
        
            
            self.utils.openFile(self.__file_location)
            self.utils.sleepUntil(3)
            assert self.word_maincode.isWordProcessExist(), "Microsoft Word WINWORD.EXE didn't start properly"            
            ''' Connect the already running Word process to winium Driver instead of opening new one.'''
            self.word_maincode.connectWordToWebDriver()
            self.logger.info("connected Word application to WebDriver")
            
            assert self.word_maincode.isWordElementExist(self.word_maincode.driver,**{'Name':'MsoDockTop'}), "Word Element is not present"
            self.logger.info("Checked the basic Word Doc elements present in Word GUI") 
 
            #TC Step 5
            self.word_maincode.killProcess()
            self.utils.sleepUntil(2)        
 
            #TC Step 6 Open Windows Explorer, give the file path where the file is saved and delete the file
            self.utils.removeFile(self.__file_location)
            self.utils.sleepUntil(5)
            
            #TC Step 7 Open Word and click file name under file location-  Error message should display 
            self.word_maincode.startMicroSoftWordApp("WINWORD.EXE")
            self.utils.sleepUntil(5)

            self.word_maincode.clickOpenOtherDocumentsButton()
            self.word_maincode.clickOpenWindowBrowseButton()
            
            self.word_maincode.enterFileNameInOpenBrowseWindow(self.__file_location)
            self.utils.sleepUntil(5)
            self.logger.info("Entered Filename Open browse window as:{}".format(self.__file_location))
            
            self.word_maincode.clickOpenButtonOfOpenBrowseWindow()
            self.logger.info("Clicked Browse window->Open button")

            self.utils.sleepUntil(5)
            assert self.word_maincode.isWordElementExist(self.word_maincode.getWordPopUpWindowElement(), **{'Name' : 'OK'}), "The Word error dialog didn't pop-up"
            self.logger.info("Word error Pop-up is displayed")
            
            self.utils.sleepUntil(5)
            self.word_maincode.isWordElementTextExist(self.word_maincode.getWordPopUpWindowElement(), "File not found" , **{'ClassName' : 'Element'} )
            self.logger.info("Word error Pop-up displayed 'File not found' message")
         
            #TC Step 8 Click OK and then close Word clicking 'X' button
            
            self.word_maincode.clickOKOnWordOpenDialog()
            self.logger.info("Clicked OK button in file not found error pop-up")
            
            self.utils.sleepUntil(5)
            self.word_maincode.clickCancelOnWordOpenDialog()
            self.logger.info("Clicked Cancel button in Open dialog pop-up")
            
            self.utils.sleepUntil(5)
            self.word_maincode.clickWordAppBackButton()
            self.logger.info("Clicked Word back button")
            
                    
            self.word_maincode.clickCloseButtonWordApp()
            self.utils.sleepUntil(3)
               
            self.word_maincode.isWordProcessExist() == False, "Microsoft Word process still running"
            return True, "", func.co_name

        except (AssertionError, AttributeError, TypeError, WebDriverException, NoSuchElementException) as e:
            return False, e, func.co_name
        finally:
            '''kill the web driver '''
            pid.terminate()
            self.logger.info("Killed the webDriver process of Word application")
            self.word_maincode.killProcess()
                             
    