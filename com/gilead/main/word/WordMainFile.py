'''
Created on Oct 11, 2018
@author: pparamasivan
'''

from definitions import WORD_PROCESS

from selenium import webdriver
from docx import Document
import pyautogui

class WordMainClass():
    __process = WORD_PROCESS

    
    def __init__(self, utilsRef, loggerRef):
        self.utils = utilsRef
        self.logger = loggerRef
        self.driver = ""
    
    def startMicroSoftWordApp(self, file):    
        self.utils.openApplication(file)

    def killProcess(self):
        self.utils.killProcess(self.__process)
        
    def createWordDocument(self, filename):
        self.doc = Document()
        self.doc.add_paragraph("This is line one of the word document")
        self.doc.add_paragraph("This is next Line in the document")
        self.doc.add_paragraph("I love Python")
        self.doc.save(filename)
        self.logger.info("Created Word Document at {}".format(filename))
 
            
    def isWordProcessExist(self):
        return self.utils.isProcessExist(self.__process)

    def getWordDocumentElement(self, nameOfDocument):
        doc_element = self.driver.find_element_by_name(nameOfDocument)
        return doc_element    
    
    def getWordDocumentContentText(self):
        parent_element = self.driver.find_element_by_class_name("OpusApp")  
        pane_element = parent_element.find_element_by_class_name("_WwF")
        child_element = pane_element.find_element_by_class_name("_WwB")
        doc_element = child_element.find_element_by_class_name("_WwG")
        page1_element = doc_element.find_element_by_name("Page 1")
        content_ele = page1_element.find_element_by_name("Page 1 content")
        return content_ele.text
  
    def clickCloseButtonOfWordApp(self):
        parent_element = self.driver.find_element_by_class_name("OpusApp")
        doc_top_element = parent_element.find_element_by_name("MsoDockTop")
        doc_top_element.find_element_by_name("Close").click()

    def connectWordToWebDriver(self):
        self.driver = webdriver.Remote(command_executor="http://localhost:9999",
                                       desired_capabilities={
                                           'app': r'C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE',
                                           'debugConnectToRunningApp': 'true',
                                           "args": "-port 345"
                                        })
        

    def isWordElementExist(self, webelement, **property1):
        return self.utils.isElementPresent( webelement,  **property1)

    def isWordElementTextExist(self, webelement, text , **property1):
        return self.utils.isElementTextPresent(webelement, text , **property1)

    def clickBlankDocument(self):
        parent_element = self.driver.find_element_by_name("Word")
        pane1 = parent_element.find_element_by_class_name("FullpageUIHost")
        pane2 = pane1.find_element_by_class_name("NetUIFullpageUIWindow")
        bs_element = pane2.find_element_by_name("Backstage view")
        group_element = bs_element.find_element_by_class_name("NetUISlabContainer")
        featured_element = group_element.find_element_by_name("Featured")
        featured_element.find_element_by_name("Blank document").submit()

    def enterTextInDocument(self, text):
        parent_element = self.driver.find_element_by_class_name("OpusApp")
        pane_element = parent_element.find_element_by_class_name("_WwF")
        child_element = pane_element.find_element_by_class_name("_WwB")
        doc_element = child_element.find_element_by_class_name("_WwG")
        page1_element = doc_element.find_element_by_name("Page 1")
        content_element=page1_element.find_element_by_name("Page 1 content")
        content_element.submit()
        content_element.send_keys(text)
        
    def clickDontSaveInMsPopUp(self):
        parent_element = self.driver.find_element_by_class_name("OpusApp")
        pane_element = parent_element.find_element_by_name("Microsoft Word")
        pane_element.find_element_by_name("Don't Save").click()
    
    def clickCloseWordAppButton(self):
        parent_element = self.driver.find_element_by_class_name("OpusApp")
        MsoDockTop_element = parent_element.find_element_by_class_name("MsoCommandBarDock")
        ribbont_element = MsoDockTop_element.find_element_by_class_name("MsoCommandBar")
        pane_element = ribbont_element.find_element_by_class_name("NetUIHWNDElement")
        ribbon_element = pane_element.find_element_by_class_name("NetUInetpane")
        ribbon_element.find_element_by_name("Close").click()

    def clickWordDontSaveMenuOption(self):
        pyautogui.hotkey('alt','f4')
    
    
    def clickOpenOtherDocumentsButton(self):
        parent_element = self.driver.find_element_by_name("Word")
        backstage_element = parent_element.find_element_by_name("Backstage view")
        backstage_element.find_element_by_name("Open Other Documents").click()
    
    def clickOpenWindowBrowseButton(self):
        parent_element = self.driver.find_element_by_name("Word")
        backstage_element = parent_element.find_element_by_name("Backstage view")
        backstage_element.find_element_by_name("Browse").click()


    def clickOpenButtonOfOpenBrowseWindow(self):
        browse_window_element = self.driver.find_element_by_class_name("#32770")
        elements =  browse_window_element.find_elements_by_class_name("Button")
        for element in elements:
            if element.get_attribute("Name") == "Open":
                element.click()


    def getWordPopUpWindowElement(self):
        parent_element = self.driver.find_element_by_name("Word")
        open_dialog_element = parent_element.find_element_by_class_name("#32770")
        return open_dialog_element
    
        
    def clickOKOnWordOpenDialog(self):
        open_dialog_element = self.getWordPopUpWindowElement()
        open_dialog_element.find_element_by_name("OK").click()
    
    def clickCancelOnWordOpenDialog(self):
        browse_window_element = self.driver.find_element_by_class_name("#32770")
        elements =  browse_window_element.find_elements_by_class_name("Button")
        for element in elements:
            if element.get_attribute("Name") == "Cancel":
                element.click()
    
    def clickWordAppBackButton(self):
        parent_element = self.driver.find_element_by_class_name("OpusApp")
        parent_element.find_element_by_name("Back").click()    
    
    def clickCloseButtonWordApp(self):
        parent_element = self.driver.find_element_by_class_name("OpusApp")
        pane_element = parent_element.find_element_by_class_name("FullpageUIHost")
        pane_child_element = pane_element.find_element_by_class_name("NetUIFullpageUIWindow")
        pane_child_element.find_element_by_name("Close").click()

    def enterFileNameInOpenBrowseWindow(self, fileName):
        browse_window_element = self.driver.find_element_by_class_name("#32770")
        browse_window_element.find_element_by_class_name("Edit").send_keys(fileName)
     